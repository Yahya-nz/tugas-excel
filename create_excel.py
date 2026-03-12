#!/usr/bin/env python3
"""
Create Excel workbook for Finance homework:
- Monthly returns for AAPL, VFIAX, S&P 500
- Risk-free rate from DGS1MO and Fama-French RF
- Excess returns
- Descriptive statistics
- Market Model (CAPM) regression with residuals
- Fama-French 3-Factor + Momentum regression
- Appraisal ratios
Also creates Word document with answers.
"""

import csv
import re
import numpy as np
from datetime import datetime
from collections import OrderedDict

import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, numbers
from openpyxl.utils import get_column_letter

# ============================================================
# 1. PARSE ALL DATA
# ============================================================

def parse_yahoo_csv(filename):
    """Parse Yahoo Finance CSV, return dict of (year, month) -> adj_close, skipping dividend rows."""
    data = {}
    with open(filename, 'r') as f:
        reader = csv.reader(f)
        header = next(reader)
        for row in reader:
            date_str = row[0]
            if 'Dividend' in date_str or 'Distribution' in date_str:
                continue
            parts = date_str.split()
            if len(parts) < 3:
                continue
            month_str = parts[0]
            year = int(parts[-1])
            months = {'Jan':1,'Feb':2,'Mar':3,'Apr':4,'May':5,'Jun':6,
                      'Jul':7,'Aug':8,'Sep':9,'Oct':10,'Nov':11,'Dec':12}
            month = months.get(month_str, 0)
            if month == 0:
                continue
            adj_close = row[5]
            if adj_close == '' or adj_close == '-':
                adj_close = row[4]
            try:
                val = float(adj_close)
            except:
                continue
            key = (year, month)
            day = int(parts[1]) if len(parts) >= 3 else 1
            if key not in data or day <= 1:
                data[key] = val
    return data


def parse_ff_factors(filename):
    """Parse Fama-French factors file, return dict of YYYYMM -> {factor: value}."""
    data = {}
    with open(filename, 'r') as f:
        lines = f.readlines()
    in_monthly = False
    for line in lines:
        line = line.rstrip()
        if 'Mkt-RF' in line and 'SMB' in line and 'HML' in line:
            in_monthly = True
            continue
        if 'Annual' in line:
            break
        if not in_monthly:
            continue
        parts = line.split()
        if len(parts) >= 5:
            try:
                yyyymm = int(parts[0])
                if yyyymm < 100000:
                    continue
                data[yyyymm] = {
                    'Mkt-RF': float(parts[1]),
                    'SMB': float(parts[2]),
                    'HML': float(parts[3]),
                    'RF': float(parts[4])
                }
            except:
                continue
    return data


def parse_momentum(filename):
    """Parse momentum factor file, return dict of YYYYMM -> Mom value."""
    data = {}
    with open(filename, 'r') as f:
        lines = f.readlines()
    in_monthly = False
    for line in lines:
        line = line.rstrip()
        if 'Mom' in line and len(line.strip().split()) <= 2:
            in_monthly = True
            continue
        if 'Annual' in line:
            break
        if not in_monthly:
            continue
        parts = line.split()
        if len(parts) >= 2:
            try:
                yyyymm = int(parts[0])
                if yyyymm < 100000:
                    continue
                data[yyyymm] = float(parts[1])
            except:
                continue
    return data


def parse_dgs1mo(filename):
    """Parse DGS1MO Excel file, return dict of (year, month) -> avg monthly rate in %."""
    wb = openpyxl.load_workbook(filename, data_only=True)
    ws = wb['Daily']
    monthly_rates = {}
    monthly_counts = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        date_val, rate = row[0], row[1]
        if date_val is None or rate is None:
            continue
        if isinstance(rate, str):
            if rate == '.' or rate == '#N/A':
                continue
            try:
                rate = float(rate)
            except:
                continue
        if isinstance(date_val, datetime):
            y, m = date_val.year, date_val.month
        else:
            continue
        key = (y, m)
        if key not in monthly_rates:
            monthly_rates[key] = 0.0
            monthly_counts[key] = 0
        monthly_rates[key] += rate
        monthly_counts[key] += 1
    result = {}
    for key in monthly_rates:
        result[key] = monthly_rates[key] / monthly_counts[key]
    return result


# Parse all data
print("Parsing data files...")
aapl = parse_yahoo_csv('AAPL_Historical_Prices.csv')
vfiax = parse_yahoo_csv('VFIAX_Historical_Prices.csv')
gspc = parse_yahoo_csv('GSPC_Historical_Prices.csv')
ff_factors = parse_ff_factors('F-F_Research_Data_Factors.txt')
momentum = parse_momentum('F-F_Momentum_Factor.txt')
dgs1mo = parse_dgs1mo('DGS1MO.xlsx')

all_months = sorted(set(aapl.keys()) & set(vfiax.keys()) & set(gspc.keys()))
print(f"Common months with price data: {len(all_months)}")
print(f"Range: {all_months[0]} to {all_months[-1]}")

return_months = all_months[1:]
print(f"Return months: {len(return_months)}")

# ============================================================
# 1b. COMPUTE RETURNS & REGRESSIONS IN PYTHON (for residuals & Word)
# ============================================================

# Compute monthly returns
aapl_ret, vfiax_ret, gspc_ret = [], [], []
rf_list, mktrf_list, smb_list, hml_list, mom_list = [], [], [], [], []
aapl_excess, vfiax_excess = [], []

for i, mk in enumerate(return_months):
    prev = all_months[i]  # return_months[i] = all_months[i+1], prev = all_months[i]
    r_aapl = aapl[mk] / aapl[prev] - 1
    r_vfiax = vfiax[mk] / vfiax[prev] - 1
    r_gspc = gspc[mk] / gspc[prev] - 1

    yyyymm = mk[0] * 100 + mk[1]
    rf = ff_factors[yyyymm]['RF'] / 100 if yyyymm in ff_factors else 0
    mktrf = ff_factors[yyyymm]['Mkt-RF'] / 100 if yyyymm in ff_factors else 0
    smb = ff_factors[yyyymm]['SMB'] / 100 if yyyymm in ff_factors else 0
    hml = ff_factors[yyyymm]['HML'] / 100 if yyyymm in ff_factors else 0
    mom_val = momentum[yyyymm] / 100 if yyyymm in momentum else 0

    aapl_ret.append(r_aapl)
    vfiax_ret.append(r_vfiax)
    gspc_ret.append(r_gspc)
    rf_list.append(rf)
    mktrf_list.append(mktrf)
    smb_list.append(smb)
    hml_list.append(hml)
    mom_list.append(mom_val)
    aapl_excess.append(r_aapl - rf)
    vfiax_excess.append(r_vfiax - rf)

# Convert to numpy
aapl_excess_np = np.array(aapl_excess)
vfiax_excess_np = np.array(vfiax_excess)
mktrf_np = np.array(mktrf_list)
smb_np = np.array(smb_list)
hml_np = np.array(hml_list)
mom_np = np.array(mom_list)
n = len(return_months)

# --- CAPM Regression ---
def ols_simple(y, x):
    X = np.column_stack([np.ones(len(x)), x])
    beta = np.linalg.lstsq(X, y, rcond=None)[0]
    y_hat = X @ beta
    resid = y - y_hat
    ss_res = np.sum(resid**2)
    ss_tot = np.sum((y - np.mean(y))**2)
    r2 = 1 - ss_res / ss_tot
    k = X.shape[1]
    se_resid = np.sqrt(ss_res / (n - k))
    cov_beta = se_resid**2 * np.linalg.inv(X.T @ X)
    se_beta = np.sqrt(np.diag(cov_beta))
    t_stats = beta / se_beta
    return {'alpha': beta[0], 'beta': beta[1], 'r2': r2,
            'se_alpha': se_beta[0], 'se_beta': se_beta[1],
            't_alpha': t_stats[0], 't_beta': t_stats[1],
            'resid': resid, 'se_resid': se_resid, 'y_hat': y_hat}

def ols_multi(y, X_vars):
    X = np.column_stack([np.ones(len(y))] + X_vars)
    beta = np.linalg.lstsq(X, y, rcond=None)[0]
    y_hat = X @ beta
    resid = y - y_hat
    ss_res = np.sum(resid**2)
    ss_tot = np.sum((y - np.mean(y))**2)
    r2 = 1 - ss_res / ss_tot
    k = X.shape[1]
    se_resid = np.sqrt(ss_res / (n - k))
    cov_beta = se_resid**2 * np.linalg.inv(X.T @ X)
    se_beta = np.sqrt(np.diag(cov_beta))
    t_stats = beta / se_beta
    f_stat = ((ss_tot - ss_res) / (k - 1)) / (ss_res / (n - k))
    return {'betas': beta, 'r2': r2, 'se': se_beta, 't_stats': t_stats,
            'resid': resid, 'se_resid': se_resid, 'f_stat': f_stat}

# CAPM
capm_aapl = ols_simple(aapl_excess_np, mktrf_np)
capm_vfiax = ols_simple(vfiax_excess_np, mktrf_np)

# FF4
ff4_aapl = ols_multi(aapl_excess_np, [mktrf_np, smb_np, hml_np, mom_np])
ff4_vfiax = ols_multi(vfiax_excess_np, [mktrf_np, smb_np, hml_np, mom_np])

# Appraisal ratios
appraisal_aapl_capm = capm_aapl['alpha'] / np.std(capm_aapl['resid'], ddof=1)
appraisal_vfiax_capm = capm_vfiax['alpha'] / np.std(capm_vfiax['resid'], ddof=1)
appraisal_aapl_ff4 = ff4_aapl['betas'][0] / np.std(ff4_aapl['resid'], ddof=1)
appraisal_vfiax_ff4 = ff4_vfiax['betas'][0] / np.std(ff4_vfiax['resid'], ddof=1)

print(f"\n=== CAPM Results ===")
print(f"AAPL:  alpha={capm_aapl['alpha']:.4f}, beta={capm_aapl['beta']:.4f}, R2={capm_aapl['r2']:.4f}, t(a)={capm_aapl['t_alpha']:.4f}, t(b)={capm_aapl['t_beta']:.4f}")
print(f"VFIAX: alpha={capm_vfiax['alpha']:.4f}, beta={capm_vfiax['beta']:.4f}, R2={capm_vfiax['r2']:.4f}, t(a)={capm_vfiax['t_alpha']:.4f}, t(b)={capm_vfiax['t_beta']:.4f}")
print(f"\n=== FF4 Results ===")
print(f"AAPL:  betas={[f'{b:.4f}' for b in ff4_aapl['betas']]}, R2={ff4_aapl['r2']:.4f}")
print(f"       t-stats={[f'{t:.4f}' for t in ff4_aapl['t_stats']]}")
print(f"VFIAX: betas={[f'{b:.4f}' for b in ff4_vfiax['betas']]}, R2={ff4_vfiax['r2']:.4f}")
print(f"       t-stats={[f'{t:.4f}' for t in ff4_vfiax['t_stats']]}")
print(f"\n=== Appraisal Ratios (CAPM) ===")
print(f"AAPL:  {appraisal_aapl_capm:.4f}")
print(f"VFIAX: {appraisal_vfiax_capm:.4f}")

# ============================================================
# 2. CREATE EXCEL WORKBOOK
# ============================================================

wb = Workbook()

# Styles
header_font = Font(bold=True, size=11)
title_font = Font(bold=True, size=14)
subtitle_font = Font(bold=True, size=12)
pct_format = '0.00%'
pct4_format = '0.0000%'
num2_format = '0.00'
num4_format = '0.0000'
thin_border = Border(
    left=Side(style='thin'), right=Side(style='thin'),
    top=Side(style='thin'), bottom=Side(style='thin')
)
header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
header_font_white = Font(bold=True, size=11, color='FFFFFF')
light_fill = PatternFill(start_color='D6E4F0', end_color='D6E4F0', fill_type='solid')

def style_header(ws, row, cols, fill=None, font=None):
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = font or header_font_white
        cell.fill = fill or header_fill
        cell.alignment = Alignment(horizontal='center', wrap_text=True)
        cell.border = thin_border

def auto_width(ws):
    for col in ws.columns:
        max_length = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max_length + 3, 22)


# ============================================================
# SHEET 1: Monthly Data & Returns
# ============================================================
ws1 = wb.active
ws1.title = "Monthly Returns"

# Headers - added residual columns (Q=AAPL CAPM resid, R=VFIAX CAPM resid)
headers1 = ['Date', 'AAPL Price', 'VFIAX NAV', 'S&P 500',
            'AAPL Return', 'VFIAX Return', 'S&P 500 Return',
            'RF (FF, %/mo)', 'RF (DGS1MO, %/yr)',
            'AAPL Excess', 'VFIAX Excess', 'S&P 500 Excess',
            'Mkt-RF (FF)', 'SMB (FF)', 'HML (FF)', 'Mom',
            'AAPL CAPM Resid', 'VFIAX CAPM Resid']

ws1.cell(row=1, column=1, value="Monthly Price Data, Returns & Fama-French Factors").font = title_font
ws1.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers1))

row = 3
for i, h in enumerate(headers1, 1):
    ws1.cell(row=row, column=i, value=h)
style_header(ws1, row, len(headers1))

# Data rows
data_start_row = 4

for idx, month_key in enumerate(all_months):
    r = data_start_row + idx
    y, m = month_key

    # Date
    ws1.cell(row=r, column=1, value=datetime(y, m, 1))
    ws1.cell(row=r, column=1).number_format = 'MMM YYYY'
    ws1.cell(row=r, column=1).border = thin_border

    # Prices
    ws1.cell(row=r, column=2, value=aapl[month_key])
    ws1.cell(row=r, column=2).number_format = '#,##0.00'
    ws1.cell(row=r, column=2).border = thin_border

    ws1.cell(row=r, column=3, value=vfiax[month_key])
    ws1.cell(row=r, column=3).number_format = '#,##0.00'
    ws1.cell(row=r, column=3).border = thin_border

    ws1.cell(row=r, column=4, value=gspc[month_key])
    ws1.cell(row=r, column=4).number_format = '#,##0.00'
    ws1.cell(row=r, column=4).border = thin_border

    # Returns (from 2nd month onward)
    if idx > 0:
        prev_r = r - 1
        ws1.cell(row=r, column=5).value = f'=B{r}/B{prev_r}-1'
        ws1.cell(row=r, column=5).number_format = pct_format
        ws1.cell(row=r, column=5).border = thin_border

        ws1.cell(row=r, column=6).value = f'=C{r}/C{prev_r}-1'
        ws1.cell(row=r, column=6).number_format = pct_format
        ws1.cell(row=r, column=6).border = thin_border

        ws1.cell(row=r, column=7).value = f'=D{r}/D{prev_r}-1'
        ws1.cell(row=r, column=7).number_format = pct_format
        ws1.cell(row=r, column=7).border = thin_border

    # Fama-French RF
    yyyymm = y * 100 + m
    if yyyymm in ff_factors:
        ws1.cell(row=r, column=8, value=ff_factors[yyyymm]['RF'] / 100)
        ws1.cell(row=r, column=8).number_format = pct4_format
        ws1.cell(row=r, column=8).border = thin_border

    # DGS1MO
    if month_key in dgs1mo:
        ws1.cell(row=r, column=9, value=dgs1mo[month_key] / 100)
        ws1.cell(row=r, column=9).number_format = pct4_format
        ws1.cell(row=r, column=9).border = thin_border

    # Excess returns
    if idx > 0:
        ws1.cell(row=r, column=10).value = f'=E{r}-H{r}'
        ws1.cell(row=r, column=10).number_format = pct_format
        ws1.cell(row=r, column=10).border = thin_border

        ws1.cell(row=r, column=11).value = f'=F{r}-H{r}'
        ws1.cell(row=r, column=11).number_format = pct_format
        ws1.cell(row=r, column=11).border = thin_border

        ws1.cell(row=r, column=12).value = f'=G{r}-H{r}'
        ws1.cell(row=r, column=12).number_format = pct_format
        ws1.cell(row=r, column=12).border = thin_border

    # FF Factors
    if yyyymm in ff_factors:
        ws1.cell(row=r, column=13, value=ff_factors[yyyymm]['Mkt-RF'] / 100)
        ws1.cell(row=r, column=13).number_format = pct_format
        ws1.cell(row=r, column=13).border = thin_border

        ws1.cell(row=r, column=14, value=ff_factors[yyyymm]['SMB'] / 100)
        ws1.cell(row=r, column=14).number_format = pct_format
        ws1.cell(row=r, column=14).border = thin_border

        ws1.cell(row=r, column=15, value=ff_factors[yyyymm]['HML'] / 100)
        ws1.cell(row=r, column=15).number_format = pct_format
        ws1.cell(row=r, column=15).border = thin_border

    # Momentum
    if yyyymm in momentum:
        ws1.cell(row=r, column=16, value=momentum[yyyymm] / 100)
        ws1.cell(row=r, column=16).number_format = pct_format
        ws1.cell(row=r, column=16).border = thin_border

    # CAPM Residuals (from computed regression)
    if idx > 0:
        ri = idx - 1  # index in return arrays
        ws1.cell(row=r, column=17, value=capm_aapl['resid'][ri])
        ws1.cell(row=r, column=17).number_format = pct_format
        ws1.cell(row=r, column=17).border = thin_border

        ws1.cell(row=r, column=18, value=capm_vfiax['resid'][ri])
        ws1.cell(row=r, column=18).number_format = pct_format
        ws1.cell(row=r, column=18).border = thin_border

last_data_row = data_start_row + len(all_months) - 1
first_return_row = data_start_row + 1

auto_width(ws1)

# ============================================================
# SHEET 2: Descriptive Statistics
# ============================================================
ws2 = wb.create_sheet("Descriptive Statistics")
ws2.cell(row=1, column=1, value="Descriptive Statistics of Monthly Returns").font = title_font
ws2.merge_cells(start_row=1, start_column=1, end_row=1, end_column=7)

stats_headers = ['Statistic', 'AAPL', 'VFIAX', 'S&P 500', 'AAPL Excess', 'VFIAX Excess', 'S&P 500 Excess']
for i, h in enumerate(stats_headers, 1):
    ws2.cell(row=3, column=i, value=h)
style_header(ws2, 3, len(stats_headers))

ret_cols = ['E', 'F', 'G', 'J', 'K', 'L']
ret_range = [f"'Monthly Returns'!{c}{first_return_row}:{c}{last_data_row}" for c in ret_cols]

stats = [
    ('Mean', 'AVERAGE'),
    ('Median', 'MEDIAN'),
    ('Std Dev', 'STDEV.S'),
    ('Variance', 'VAR.S'),
    ('Min', 'MIN'),
    ('Max', 'MAX'),
    ('Skewness', 'SKEW'),
    ('Kurtosis', 'KURT'),
    ('Count', 'COUNT'),
]

for i, (label, func) in enumerate(stats):
    r = 4 + i
    ws2.cell(row=r, column=1, value=label)
    ws2.cell(row=r, column=1).font = header_font
    ws2.cell(row=r, column=1).border = thin_border
    for j, rng in enumerate(ret_range):
        ws2.cell(row=r, column=j+2).value = f'={func}({rng})'
        if label == 'Count':
            ws2.cell(row=r, column=j+2).number_format = '0'
        elif label in ('Skewness', 'Kurtosis'):
            ws2.cell(row=r, column=j+2).number_format = num4_format
        else:
            ws2.cell(row=r, column=j+2).number_format = pct4_format
        ws2.cell(row=r, column=j+2).border = thin_border

# Annualized stats
ann_row = 4 + len(stats) + 1
ws2.cell(row=ann_row, column=1, value="Annualized Statistics").font = subtitle_font
ws2.merge_cells(start_row=ann_row, start_column=1, end_row=ann_row, end_column=7)

ann_headers = ['Statistic', 'AAPL', 'VFIAX', 'S&P 500']
for i, h in enumerate(ann_headers, 1):
    ws2.cell(row=ann_row+1, column=i, value=h)
style_header(ws2, ann_row+1, len(ann_headers))

mean_row = 4
std_row = 6

ws2.cell(row=ann_row+2, column=1, value='Annualized Return')
ws2.cell(row=ann_row+2, column=1).font = header_font
ws2.cell(row=ann_row+2, column=1).border = thin_border
for j in range(3):
    col = j + 2
    ws2.cell(row=ann_row+2, column=col).value = f'=(1+{get_column_letter(col)}{mean_row})^12-1'
    ws2.cell(row=ann_row+2, column=col).number_format = pct_format
    ws2.cell(row=ann_row+2, column=col).border = thin_border

ws2.cell(row=ann_row+3, column=1, value='Annualized Std Dev')
ws2.cell(row=ann_row+3, column=1).font = header_font
ws2.cell(row=ann_row+3, column=1).border = thin_border
for j in range(3):
    col = j + 2
    ws2.cell(row=ann_row+3, column=col).value = f'={get_column_letter(col)}{std_row}*SQRT(12)'
    ws2.cell(row=ann_row+3, column=col).number_format = pct_format
    ws2.cell(row=ann_row+3, column=col).border = thin_border

ws2.cell(row=ann_row+4, column=1, value='Sharpe Ratio')
ws2.cell(row=ann_row+4, column=1).font = header_font
ws2.cell(row=ann_row+4, column=1).border = thin_border
for j in range(3):
    col = j + 2
    excess_col = j + 5
    ws2.cell(row=ann_row+4, column=col).value = f'=((1+{get_column_letter(excess_col)}{mean_row})^12-1)/({get_column_letter(col)}{std_row}*SQRT(12))'
    ws2.cell(row=ann_row+4, column=col).number_format = num4_format
    ws2.cell(row=ann_row+4, column=col).border = thin_border

# Correlation matrix
corr_row = ann_row + 7
ws2.cell(row=corr_row, column=1, value="Correlation Matrix").font = subtitle_font
ws2.merge_cells(start_row=corr_row, start_column=1, end_row=corr_row, end_column=4)

corr_headers = ['', 'AAPL', 'VFIAX', 'S&P 500']
for i, h in enumerate(corr_headers, 1):
    ws2.cell(row=corr_row+1, column=i, value=h)
style_header(ws2, corr_row+1, len(corr_headers))

assets = ['AAPL', 'VFIAX', 'S&P 500']
ret_cols_corr = ['E', 'F', 'G']

for i in range(3):
    r = corr_row + 2 + i
    ws2.cell(row=r, column=1, value=assets[i])
    ws2.cell(row=r, column=1).font = header_font
    ws2.cell(row=r, column=1).border = thin_border
    for j in range(3):
        rng1 = f"'Monthly Returns'!{ret_cols_corr[i]}{first_return_row}:{ret_cols_corr[i]}{last_data_row}"
        rng2 = f"'Monthly Returns'!{ret_cols_corr[j]}{first_return_row}:{ret_cols_corr[j]}{last_data_row}"
        ws2.cell(row=r, column=j+2).value = f'=CORREL({rng1},{rng2})'
        ws2.cell(row=r, column=j+2).number_format = num4_format
        ws2.cell(row=r, column=j+2).border = thin_border

auto_width(ws2)

# ============================================================
# SHEET 3: Market Model (CAPM) Regression
# ============================================================
ws3 = wb.create_sheet("CAPM Regression")
ws3.cell(row=1, column=1, value="Market Model (CAPM) Regression").font = title_font
ws3.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)

ws3.cell(row=2, column=1, value="R_i - R_f = \u03b1 + \u03b2(R_m - R_f) + \u03b5").font = Font(italic=True, size=11)

y_aapl_rng = f"'Monthly Returns'!J{first_return_row}:J{last_data_row}"
y_vfiax_rng = f"'Monthly Returns'!K{first_return_row}:K{last_data_row}"
x_mkt_rng = f"'Monthly Returns'!M{first_return_row}:M{last_data_row}"

# AAPL CAPM
ws3.cell(row=4, column=1, value="AAPL Market Model").font = subtitle_font
capm_headers = ['Statistic', 'Value']
for i, h in enumerate(capm_headers, 1):
    ws3.cell(row=5, column=i, value=h)
style_header(ws3, 5, 2)

capm_stats_aapl = [
    ('Alpha (\u03b1)', f'=INDEX(LINEST({y_aapl_rng},{x_mkt_rng},TRUE,TRUE),1,2)'),
    ('Beta (\u03b2)', f'=INDEX(LINEST({y_aapl_rng},{x_mkt_rng},TRUE,TRUE),1,1)'),
    ('R-squared', f'=RSQ({y_aapl_rng},{x_mkt_rng})'),
    ('Std Error (\u03b1)', f'=INDEX(LINEST({y_aapl_rng},{x_mkt_rng},TRUE,TRUE),2,2)'),
    ('Std Error (\u03b2)', f'=INDEX(LINEST({y_aapl_rng},{x_mkt_rng},TRUE,TRUE),2,1)'),
    ('t-stat (\u03b1)', '=B6/B9'),
    ('t-stat (\u03b2)', '=B7/B10'),
    ('F-statistic', f'=INDEX(LINEST({y_aapl_rng},{x_mkt_rng},TRUE,TRUE),4,1)'),
    ('Observations', f'=COUNT({y_aapl_rng})'),
    ('Std Dev of Residuals', f'=STDEV.S(\'Monthly Returns\'!Q{first_return_row}:Q{last_data_row})'),
    ('Appraisal Ratio', '=B6/B15'),
]

for i, (label, formula) in enumerate(capm_stats_aapl):
    r = 6 + i
    ws3.cell(row=r, column=1, value=label)
    ws3.cell(row=r, column=1).font = header_font
    ws3.cell(row=r, column=1).border = thin_border
    ws3.cell(row=r, column=2).value = formula
    ws3.cell(row=r, column=2).number_format = num4_format
    ws3.cell(row=r, column=2).border = thin_border

# VFIAX CAPM
vfiax_start = 19
ws3.cell(row=vfiax_start, column=1, value="VFIAX Market Model").font = subtitle_font

for i, h in enumerate(capm_headers, 1):
    ws3.cell(row=vfiax_start+1, column=i, value=h)
style_header(ws3, vfiax_start+1, 2)

capm_stats_vfiax = [
    ('Alpha (\u03b1)', f'=INDEX(LINEST({y_vfiax_rng},{x_mkt_rng},TRUE,TRUE),1,2)'),
    ('Beta (\u03b2)', f'=INDEX(LINEST({y_vfiax_rng},{x_mkt_rng},TRUE,TRUE),1,1)'),
    ('R-squared', f'=RSQ({y_vfiax_rng},{x_mkt_rng})'),
    ('Std Error (\u03b1)', f'=INDEX(LINEST({y_vfiax_rng},{x_mkt_rng},TRUE,TRUE),2,2)'),
    ('Std Error (\u03b2)', f'=INDEX(LINEST({y_vfiax_rng},{x_mkt_rng},TRUE,TRUE),2,1)'),
    ('t-stat (\u03b1)', f'=B{vfiax_start+2}/B{vfiax_start+5}'),
    ('t-stat (\u03b2)', f'=B{vfiax_start+3}/B{vfiax_start+6}'),
    ('F-statistic', f'=INDEX(LINEST({y_vfiax_rng},{x_mkt_rng},TRUE,TRUE),4,1)'),
    ('Observations', f'=COUNT({y_vfiax_rng})'),
    ('Std Dev of Residuals', f'=STDEV.S(\'Monthly Returns\'!R{first_return_row}:R{last_data_row})'),
    ('Appraisal Ratio', f'=B{vfiax_start+2}/B{vfiax_start+11}'),
]

for i, (label, formula) in enumerate(capm_stats_vfiax):
    r = vfiax_start + 2 + i
    ws3.cell(row=r, column=1, value=label)
    ws3.cell(row=r, column=1).font = header_font
    ws3.cell(row=r, column=1).border = thin_border
    ws3.cell(row=r, column=2).value = formula
    ws3.cell(row=r, column=2).number_format = num4_format
    ws3.cell(row=r, column=2).border = thin_border

auto_width(ws3)

# ============================================================
# SHEET 4: Fama-French 4-Factor
# ============================================================
ws4 = wb.create_sheet("FF4 Regression")
ws4.cell(row=1, column=1, value="Fama-French 3-Factor + Momentum (Carhart 4-Factor) Regression").font = title_font
ws4.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)

ws4.cell(row=2, column=1, value="R_i - R_f = \u03b1 + \u03b2\u2081(Mkt-RF) + \u03b2\u2082(SMB) + \u03b2\u2083(HML) + \u03b2\u2084(Mom) + \u03b5").font = Font(italic=True, size=11)

x_all_rng = f"'Monthly Returns'!M{first_return_row}:P{last_data_row}"

# AAPL FF4
ws4.cell(row=4, column=1, value="AAPL Fama-French 4-Factor").font = subtitle_font

ff4_headers = ['Statistic', 'Value']
for i, h in enumerate(ff4_headers, 1):
    ws4.cell(row=5, column=i, value=h)
style_header(ws4, 5, 2)

ff4_stats_aapl = [
    ('Alpha (\u03b1)',     f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),1,5)'),
    ('\u03b2 (Mkt-RF)',    f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),1,4)'),
    ('\u03b2 (SMB)',       f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),1,3)'),
    ('\u03b2 (HML)',       f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),1,2)'),
    ('\u03b2 (Mom)',       f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),1,1)'),
    ('R-squared',     f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),3,1)'),
    ('SE (\u03b1)',        f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),2,5)'),
    ('SE (Mkt-RF)',   f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),2,4)'),
    ('SE (SMB)',      f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),2,3)'),
    ('SE (HML)',      f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),2,2)'),
    ('SE (Mom)',      f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),2,1)'),
    ('t-stat (\u03b1)',    '=B6/B12'),
    ('t-stat (Mkt-RF)', '=B7/B13'),
    ('t-stat (SMB)',  '=B8/B14'),
    ('t-stat (HML)',  '=B9/B15'),
    ('t-stat (Mom)',  '=B10/B16'),
    ('F-statistic',   f'=INDEX(LINEST({y_aapl_rng},{x_all_rng},TRUE,TRUE),4,1)'),
    ('Observations',  f'=COUNT({y_aapl_rng})'),
]

for i, (label, formula) in enumerate(ff4_stats_aapl):
    r = 6 + i
    ws4.cell(row=r, column=1, value=label)
    ws4.cell(row=r, column=1).font = header_font
    ws4.cell(row=r, column=1).border = thin_border
    ws4.cell(row=r, column=2).value = formula
    ws4.cell(row=r, column=2).number_format = num4_format
    ws4.cell(row=r, column=2).border = thin_border

# VFIAX FF4
vf_start = 26
ws4.cell(row=vf_start, column=1, value="VFIAX Fama-French 4-Factor").font = subtitle_font

for i, h in enumerate(ff4_headers, 1):
    ws4.cell(row=vf_start+1, column=i, value=h)
style_header(ws4, vf_start+1, 2)

ff4_stats_vfiax = [
    ('Alpha (\u03b1)',     f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),1,5)'),
    ('\u03b2 (Mkt-RF)',    f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),1,4)'),
    ('\u03b2 (SMB)',       f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),1,3)'),
    ('\u03b2 (HML)',       f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),1,2)'),
    ('\u03b2 (Mom)',       f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),1,1)'),
    ('R-squared',     f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),3,1)'),
    ('SE (\u03b1)',        f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),2,5)'),
    ('SE (Mkt-RF)',   f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),2,4)'),
    ('SE (SMB)',      f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),2,3)'),
    ('SE (HML)',      f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),2,2)'),
    ('SE (Mom)',      f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),2,1)'),
    ('t-stat (\u03b1)',    f'=B{vf_start+2}/B{vf_start+8}'),
    ('t-stat (Mkt-RF)', f'=B{vf_start+3}/B{vf_start+9}'),
    ('t-stat (SMB)',  f'=B{vf_start+4}/B{vf_start+10}'),
    ('t-stat (HML)',  f'=B{vf_start+5}/B{vf_start+11}'),
    ('t-stat (Mom)',  f'=B{vf_start+6}/B{vf_start+12}'),
    ('F-statistic',   f'=INDEX(LINEST({y_vfiax_rng},{x_all_rng},TRUE,TRUE),4,1)'),
    ('Observations',  f'=COUNT({y_vfiax_rng})'),
]

for i, (label, formula) in enumerate(ff4_stats_vfiax):
    r = vf_start + 2 + i
    ws4.cell(row=r, column=1, value=label)
    ws4.cell(row=r, column=1).font = header_font
    ws4.cell(row=r, column=1).border = thin_border
    ws4.cell(row=r, column=2).value = formula
    ws4.cell(row=r, column=2).number_format = num4_format
    ws4.cell(row=r, column=2).border = thin_border

auto_width(ws4)

# ============================================================
# SAVE EXCEL
# ============================================================
output_file = 'Finance_Homework.xlsx'
wb.save(output_file)
print(f"\nExcel file saved: {output_file}")
print(f"Sheets: {wb.sheetnames}")

# ============================================================
# 3. CREATE WORD DOCUMENT WITH ANSWERS
# ============================================================
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Finance Homework Answers', level=0)
doc.add_paragraph(f'Individual Stock: AAPL (Apple Inc.)')
doc.add_paragraph(f'Mutual Fund: VFIAX (Vanguard 500 Index Fund)')
doc.add_paragraph(f'Market Index: S&P 500 (^GSPC)')
doc.add_paragraph(f'Data Period: {return_months[0][0]}/{return_months[0][1]:02d} - {return_months[-1][0]}/{return_months[-1][1]:02d} ({n} months)')
doc.add_paragraph('')

# =============================================
# CHAPTER 11 ANSWERS
# =============================================
doc.add_heading('Chapter 11: Excel Master It! Problem', level=1)

# --- Part a ---
doc.add_heading('Part a: Market Model Regression', level=2)

# Summary table
doc.add_paragraph('Regression Results Summary:', style='Intense Quote')

table = doc.add_table(rows=8, cols=3, style='Light Shading Accent 1')
table.cell(0, 0).text = 'Statistic'
table.cell(0, 1).text = 'AAPL'
table.cell(0, 2).text = 'VFIAX'
stats_data = [
    ('Alpha (α)', f'{capm_aapl["alpha"]:.4f}', f'{capm_vfiax["alpha"]:.4f}'),
    ('Beta (β)', f'{capm_aapl["beta"]:.4f}', f'{capm_vfiax["beta"]:.4f}'),
    ('R-squared', f'{capm_aapl["r2"]:.4f}', f'{capm_vfiax["r2"]:.4f}'),
    ('SE(α)', f'{capm_aapl["se_alpha"]:.4f}', f'{capm_vfiax["se_alpha"]:.4f}'),
    ('SE(β)', f'{capm_aapl["se_beta"]:.4f}', f'{capm_vfiax["se_beta"]:.4f}'),
    ('t-stat(α)', f'{capm_aapl["t_alpha"]:.4f}', f'{capm_vfiax["t_alpha"]:.4f}'),
    ('t-stat(β)', f'{capm_aapl["t_beta"]:.4f}', f'{capm_vfiax["t_beta"]:.4f}'),
]
for i, (stat, v1, v2) in enumerate(stats_data):
    table.cell(i+1, 0).text = stat
    table.cell(i+1, 1).text = v1
    table.cell(i+1, 2).text = v2

doc.add_paragraph('')

# a.1
doc.add_heading('a.1) Are alpha and beta statistically different from zero?', level=3)

# Determine significance
def sig_text(t_val, name):
    abs_t = abs(t_val)
    if abs_t > 2.66:
        return f'{name} (t = {t_val:.4f}) is statistically significant at the 1% level (|t| > 2.66).'
    elif abs_t > 2.00:
        return f'{name} (t = {t_val:.4f}) is statistically significant at the 5% level (|t| > 2.00).'
    elif abs_t > 1.67:
        return f'{name} (t = {t_val:.4f}) is statistically significant at the 10% level (|t| > 1.67).'
    else:
        return f'{name} (t = {t_val:.4f}) is NOT statistically significant at conventional levels (|t| < 1.67).'

doc.add_paragraph('AAPL:', style='List Bullet')
doc.add_paragraph(sig_text(capm_aapl['t_alpha'], 'Alpha'))
doc.add_paragraph(sig_text(capm_aapl['t_beta'], 'Beta'))
doc.add_paragraph('')
doc.add_paragraph('VFIAX:', style='List Bullet')
doc.add_paragraph(sig_text(capm_vfiax['t_alpha'], 'Alpha'))
doc.add_paragraph(sig_text(capm_vfiax['t_beta'], 'Beta'))

doc.add_paragraph('')
doc.add_paragraph(
    'For both AAPL and VFIAX, the beta coefficient is highly significant, confirming that both assets '
    'have systematic risk exposure to the market. The alpha for AAPL is likely not statistically significant, '
    'meaning AAPL did not generate a statistically reliable excess return beyond what CAPM predicts. '
    'For VFIAX (an index fund), we would expect alpha to be close to zero and not significant, as index funds '
    'are designed to track the market, not to generate alpha.'
)

# a.2
doc.add_heading('a.2) Interpretation of alpha and beta', level=3)
doc.add_paragraph(
    f'AAPL Alpha = {capm_aapl["alpha"]:.4f}: This represents the monthly excess return of AAPL '
    f'beyond what is predicted by the market model. A {"positive" if capm_aapl["alpha"] > 0 else "negative"} '
    f'alpha suggests that AAPL {"outperformed" if capm_aapl["alpha"] > 0 else "underperformed"} '
    f'the CAPM prediction over this period, though this may not be statistically significant.'
)
doc.add_paragraph(
    f'AAPL Beta = {capm_aapl["beta"]:.4f}: This measures AAPL\'s systematic risk. '
    f'A beta {"greater" if capm_aapl["beta"] > 1 else "less"} than 1.0 indicates that AAPL is '
    f'{"more" if capm_aapl["beta"] > 1 else "less"} volatile than the market. '
    f'For every 1% change in the market excess return, AAPL\'s excess return is expected to change '
    f'by approximately {capm_aapl["beta"]:.2f}%.'
)
doc.add_paragraph(
    f'VFIAX Alpha = {capm_vfiax["alpha"]:.4f}: As an S&P 500 index fund, VFIAX\'s alpha is very close '
    f'to zero, which is expected. The small {"positive" if capm_vfiax["alpha"] > 0 else "negative"} value '
    f'reflects tracking error and expense ratio differences.'
)
doc.add_paragraph(
    f'VFIAX Beta = {capm_vfiax["beta"]:.4f}: Being an index fund that tracks the S&P 500, '
    f'VFIAX\'s beta is very close to 1.0, confirming it closely mirrors market movements. '
    f'This is exactly what we expect from a passively managed index fund.'
)

# a.3
doc.add_heading('a.3) Which R-squared is highest? Expected?', level=3)
higher_r2 = 'VFIAX' if capm_vfiax['r2'] > capm_aapl['r2'] else 'AAPL'
doc.add_paragraph(
    f'AAPL R² = {capm_aapl["r2"]:.4f} ({capm_aapl["r2"]*100:.2f}%)\n'
    f'VFIAX R² = {capm_vfiax["r2"]:.4f} ({capm_vfiax["r2"]*100:.2f}%)'
)
doc.add_paragraph(
    f'{higher_r2} has the higher R-squared. This is expected because:'
)
if higher_r2 == 'VFIAX':
    doc.add_paragraph(
        'VFIAX is an S&P 500 index fund that is designed to replicate the market index. Therefore, '
        'nearly all of its return variation is explained by market movements (systematic risk), '
        'resulting in a very high R-squared close to 1.0.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'AAPL, as an individual stock, has both systematic risk (captured by beta) AND firm-specific '
        '(unsystematic/idiosyncratic) risk. The unsystematic risk component is not explained by the '
        'market model, leading to a lower R-squared. This is consistent with finance theory: '
        'individual stocks carry diversifiable risk that a well-diversified portfolio (like VFIAX) '
        'eliminates.',
        style='List Bullet'
    )

# --- Part b ---
doc.add_heading('Part b: Residuals and Appraisal Ratio', level=2)

# b.1
doc.add_heading('b.1) What does the appraisal ratio measure?', level=3)
doc.add_paragraph(
    'The appraisal ratio (also known as the information ratio in some contexts) measures the '
    'risk-adjusted excess return generated by an asset beyond what the market model predicts. '
    'It is calculated as:'
)
doc.add_paragraph('Appraisal Ratio = Alpha / Standard Deviation of Residuals', style='Intense Quote')
doc.add_paragraph(
    'The numerator (alpha) captures the abnormal return — the return not explained by the market. '
    'The denominator (standard deviation of residuals) captures the unsystematic risk — the firm-specific '
    'volatility not related to market movements. Thus, the appraisal ratio tells us how much abnormal '
    'return an asset earns per unit of unsystematic risk taken. A higher appraisal ratio indicates '
    'better performance on a risk-adjusted basis, specifically adjusting for diversifiable risk.'
)

# b.2
doc.add_heading('b.2) Appraisal ratios for AAPL and VFIAX', level=3)
doc.add_paragraph(
    f'AAPL Appraisal Ratio = {appraisal_aapl_capm:.4f}\n'
    f'VFIAX Appraisal Ratio = {appraisal_vfiax_capm:.4f}'
)
better_ar = 'AAPL' if abs(appraisal_aapl_capm) > abs(appraisal_vfiax_capm) else 'VFIAX'
doc.add_paragraph(
    f'In absolute terms, {better_ar} has the {"better" if (appraisal_aapl_capm > appraisal_vfiax_capm and better_ar == "AAPL") or (appraisal_vfiax_capm > appraisal_aapl_capm and better_ar == "VFIAX") else "larger"} '
    f'appraisal ratio. VFIAX, as an index fund, is expected to have an appraisal ratio very close to zero '
    f'because its alpha should be approximately zero. Any non-zero appraisal ratio for VFIAX reflects '
    f'minor tracking errors relative to the Fama-French market factor.'
)

# b.3
doc.add_heading('b.3) Why is the appraisal ratio used more for mutual funds?', level=3)
doc.add_paragraph(
    'The appraisal ratio is used more often for mutual funds than for individual stocks for several reasons:'
)
doc.add_paragraph(
    '1. Mutual funds are managed portfolios, and the appraisal ratio directly evaluates '
    'a manager\'s skill. Alpha represents the value added by the manager\'s stock-picking ability, '
    'while the residual standard deviation represents the idiosyncratic risk the manager chose to take. '
    'The ratio thus measures the manager\'s ability to generate abnormal returns per unit of active risk.',
    style='List Bullet'
)
doc.add_paragraph(
    '2. For individual stocks, the unsystematic risk is inherent to the firm and cannot be reduced '
    'by the "stock itself." But mutual fund managers actively choose to deviate from the benchmark, '
    'so the unsystematic risk is a deliberate choice. The appraisal ratio evaluates whether that '
    'deliberate deviation was worthwhile.',
    style='List Bullet'
)
doc.add_paragraph(
    '3. In a well-diversified portfolio, unsystematic risk is largely eliminated. Mutual funds '
    'hold many stocks, so their residual risk is relatively small. The appraisal ratio is more '
    'meaningful when the residual risk is the result of active management decisions rather than '
    'simply being an unavoidable feature of holding a single security.',
    style='List Bullet'
)

# =============================================
# CHAPTER 12 ANSWERS
# =============================================
doc.add_heading('Chapter 12: Excel Master It! Problem', level=1)
doc.add_heading('Carhart Four-Factor Model', level=2)

# Summary table
doc.add_paragraph('Four-Factor Regression Results Summary:', style='Intense Quote')

table2 = doc.add_table(rows=7, cols=3, style='Light Shading Accent 1')
table2.cell(0, 0).text = 'Statistic'
table2.cell(0, 1).text = 'AAPL'
table2.cell(0, 2).text = 'VFIAX'
# betas: [alpha, Mkt-RF, SMB, HML, Mom]
ff4_data = [
    ('Alpha (α)', f'{ff4_aapl["betas"][0]:.4f}', f'{ff4_vfiax["betas"][0]:.4f}'),
    ('β(Mkt-RF)', f'{ff4_aapl["betas"][1]:.4f}', f'{ff4_vfiax["betas"][1]:.4f}'),
    ('β(SMB)', f'{ff4_aapl["betas"][2]:.4f}', f'{ff4_vfiax["betas"][2]:.4f}'),
    ('β(HML)', f'{ff4_aapl["betas"][3]:.4f}', f'{ff4_vfiax["betas"][3]:.4f}'),
    ('β(Mom)', f'{ff4_aapl["betas"][4]:.4f}', f'{ff4_vfiax["betas"][4]:.4f}'),
    ('R-squared', f'{ff4_aapl["r2"]:.4f}', f'{ff4_vfiax["r2"]:.4f}'),
]
for i, (stat, v1, v2) in enumerate(ff4_data):
    table2.cell(i+1, 0).text = stat
    table2.cell(i+1, 1).text = v1
    table2.cell(i+1, 2).text = v2

doc.add_paragraph('')

# 1.a
doc.add_heading('1.a) Explanatory power: 4-factor vs market model — higher or lower?', level=3)
doc.add_paragraph(
    f'AAPL: Market Model R² = {capm_aapl["r2"]:.4f}, Four-Factor R² = {ff4_aapl["r2"]:.4f}\n'
    f'VFIAX: Market Model R² = {capm_vfiax["r2"]:.4f}, Four-Factor R² = {ff4_vfiax["r2"]:.4f}'
)
doc.add_paragraph(
    'We expect the four-factor model to have a higher (or at least equal) R-squared than the market '
    'model. This is because the four-factor model includes all the information in the market model '
    '(Mkt-RF) plus three additional factors (SMB, HML, Mom). Adding more independent variables to '
    'a regression can only increase (or maintain) R-squared, never decrease it. '
    'The additional factors capture sources of systematic risk that the single market factor misses, '
    'such as size effects (SMB), value/growth effects (HML), and momentum effects (Mom). '
    'For AAPL specifically, these additional factors may capture variation in returns related to '
    'Apple\'s characteristics as a large-cap growth stock.'
)

# 1.b
doc.add_heading('1.b) Are alpha and betas statistically different from zero?', level=3)

doc.add_paragraph('AAPL Four-Factor Model:', style='List Bullet')
labels_ff4 = ['Alpha', 'β(Mkt-RF)', 'β(SMB)', 'β(HML)', 'β(Mom)']
for i, label in enumerate(labels_ff4):
    doc.add_paragraph(sig_text(ff4_aapl['t_stats'][i], label))

doc.add_paragraph('')
doc.add_paragraph('VFIAX Four-Factor Model:', style='List Bullet')
for i, label in enumerate(labels_ff4):
    doc.add_paragraph(sig_text(ff4_vfiax['t_stats'][i], label))

doc.add_paragraph('')
doc.add_paragraph(
    'For VFIAX (index fund), β(Mkt-RF) should be highly significant and close to 1.0, while '
    'the other factor loadings should be close to zero, reflecting its passive nature. '
    'For AAPL, the market beta should be significant. The significance of SMB, HML, and Mom '
    'depends on AAPL\'s specific characteristics during this period.'
)

# 1.c
doc.add_heading('1.c) Interpretation of betas for each independent variable', level=3)

doc.add_paragraph('AAPL:', style='List Bullet')
doc.add_paragraph(
    f'β(Mkt-RF) = {ff4_aapl["betas"][1]:.4f}: AAPL\'s sensitivity to the overall market. '
    f'A value {"above" if ff4_aapl["betas"][1] > 1 else "below"} 1.0 means AAPL is '
    f'{"more" if ff4_aapl["betas"][1] > 1 else "less"} sensitive to market movements than the average stock.'
)
doc.add_paragraph(
    f'β(SMB) = {ff4_aapl["betas"][2]:.4f}: AAPL\'s exposure to the size factor. '
    f'A {"negative" if ff4_aapl["betas"][2] < 0 else "positive"} value indicates AAPL behaves more like a '
    f'{"large" if ff4_aapl["betas"][2] < 0 else "small"}-cap stock, which is consistent with Apple being '
    f'one of the largest companies in the world.'
)
doc.add_paragraph(
    f'β(HML) = {ff4_aapl["betas"][3]:.4f}: AAPL\'s exposure to the value factor. '
    f'A {"negative" if ff4_aapl["betas"][3] < 0 else "positive"} value indicates AAPL behaves more like a '
    f'{"growth" if ff4_aapl["betas"][3] < 0 else "value"} stock, consistent with Apple\'s high market-to-book ratio.'
)
doc.add_paragraph(
    f'β(Mom) = {ff4_aapl["betas"][4]:.4f}: AAPL\'s exposure to momentum. '
    f'A {"positive" if ff4_aapl["betas"][4] > 0 else "negative"} loading suggests AAPL tends to '
    f'{"move with" if ff4_aapl["betas"][4] > 0 else "move against"} momentum trends in the market.'
)

doc.add_paragraph('')
doc.add_paragraph('VFIAX:', style='List Bullet')
doc.add_paragraph(
    f'β(Mkt-RF) = {ff4_vfiax["betas"][1]:.4f}: Very close to 1.0, as expected for an S&P 500 index fund. '
    f'The fund\'s returns move almost one-for-one with the market.'
)
doc.add_paragraph(
    f'β(SMB) = {ff4_vfiax["betas"][2]:.4f}: Close to zero, as the S&P 500 is a large-cap index and '
    f'does not have significant small-cap exposure.'
)
doc.add_paragraph(
    f'β(HML) = {ff4_vfiax["betas"][3]:.4f}: Close to zero, reflecting the broad market blend of '
    f'value and growth stocks in the S&P 500.'
)
doc.add_paragraph(
    f'β(Mom) = {ff4_vfiax["betas"][4]:.4f}: Close to zero, as a diversified index fund should not '
    f'have significant momentum exposure.'
)

# 1.d
doc.add_heading('1.d) Which R-squared is highest? Expected?', level=3)
higher_ff4 = 'VFIAX' if ff4_vfiax['r2'] > ff4_aapl['r2'] else 'AAPL'
doc.add_paragraph(
    f'Four-Factor Model:\n'
    f'AAPL R² = {ff4_aapl["r2"]:.4f} ({ff4_aapl["r2"]*100:.2f}%)\n'
    f'VFIAX R² = {ff4_vfiax["r2"]:.4f} ({ff4_vfiax["r2"]*100:.2f}%)'
)
doc.add_paragraph(
    f'{higher_ff4} has the highest R-squared in the four-factor model. This is consistent with '
    f'the market model results and is expected for the same reasons:'
)
doc.add_paragraph(
    'VFIAX, as an S&P 500 index fund, is a well-diversified portfolio with minimal idiosyncratic risk. '
    'Almost all of its return variation is driven by systematic factors, resulting in a very high R-squared.',
    style='List Bullet'
)
doc.add_paragraph(
    'AAPL, as a single stock, has substantial firm-specific risk (e.g., product launches, earnings '
    'surprises, management changes) that cannot be explained by any systematic factor model. '
    'Therefore, its R-squared will always be lower than that of a diversified portfolio.',
    style='List Bullet'
)
doc.add_paragraph(
    f'Comparing across models, both AAPL and VFIAX have (weakly) higher R-squared values in the '
    f'four-factor model than in the market model, confirming that the additional factors provide '
    f'incremental explanatory power.'
)

# Save Word document
word_file = 'Finance_Homework_Answers.docx'
doc.save(word_file)
print(f"Word file saved: {word_file}")
print("\nDone!")
